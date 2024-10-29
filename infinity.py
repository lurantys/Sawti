#MTA0MTMwNjM1NDIzMDQ5MzE5NA.G66bP0.8G_XuTI4xiILCB9rrOCG73oCBr9JMNRT7kFOpk

import discord
from discord.ext import commands
from discord import app_commands
from googletrans import Translator
import os
from docx import Document
import speech_recognition as sr
from gtts import gTTS
from pydub import AudioSegment

# Set up the bot and translator
TOKEN = 'DISCORD TOKEN'
GUILD_ID = 1266372344612655116  # Replace with your server ID for faster testing
intents = discord.Intents.default()
intents.message_content = True  # Needed for message translation in conversation mode
bot = commands.Bot(command_prefix="!", intents=intents)
translator = Translator()

# Define language choices
LANGUAGE_CHOICES = [
    app_commands.Choice(name="English", value="en"),
    app_commands.Choice(name="French", value="fr"),
    app_commands.Choice(name="Spanish", value="es"),
    app_commands.Choice(name="German", value="de"),
    app_commands.Choice(name="Italian", value="it"),
    app_commands.Choice(name="Chinese (Simplified)", value="zh-cn"),
    app_commands.Choice(name="Japanese", value="ja"),
    app_commands.Choice(name="Korean", value="ko"),
    app_commands.Choice(name="Arabic", value="ar"),
    app_commands.Choice(name="Portuguese", value="pt"),
    app_commands.Choice(name="Russian", value="ru"),
    # Add more languages as needed
]

conversation_channels = {}

@bot.event
async def on_ready():
    print(f'Logged in as {bot.user}')
    try:
        # Sync global commands only
        synced = await bot.tree.sync()  # Syncs all commands globally
        print(f"Synced {len(synced)} global commands.")
        
    except Exception as e:
        print(f"Failed to sync commands: {e}")



# Define a slash command for text translation
@bot.tree.command(name="translate", description="Translate text from one language to another")
@app_commands.describe(text="Text to translate")
@app_commands.choices(
    source_lang=LANGUAGE_CHOICES,
    target_lang=LANGUAGE_CHOICES
)
async def translate(
    interaction: discord.Interaction,
    source_lang: app_commands.Choice[str],
    target_lang: app_commands.Choice[str],
    text: str
):
    try:
        # Translate the text
        translation = translator.translate(text, src=source_lang.value, dest=target_lang.value)
        await interaction.response.send_message(f"**Translated:** {translation.text}")
    except Exception as e:
        print(f"Error: {e}")
        await interaction.response.send_message("Sorry, I couldn't translate that text.")

# Command to start conversation mode
@bot.tree.command(name="start_conversation", description="Start conversation mode to auto-translate messages")
@app_commands.choices(
    source_lang=LANGUAGE_CHOICES,
    target_lang=LANGUAGE_CHOICES
)
async def start_conversation(interaction: discord.Interaction, source_lang: app_commands.Choice[str], target_lang: app_commands.Choice[str]):
    conversation_channels[interaction.channel_id] = (source_lang.value, target_lang.value)
    await interaction.response.send_message(f"Conversation mode started between {source_lang.name} and {target_lang.name}. Type messages and I’ll translate them!")

# Command to stop conversation mode
@bot.tree.command(name="stop_conversation", description="Stop conversation mode")
async def stop_conversation(interaction: discord.Interaction):
    if interaction.channel_id in conversation_channels:
        del conversation_channels[interaction.channel_id]
        await interaction.response.send_message("Conversation mode stopped.")
    else:
        await interaction.response.send_message("Conversation mode is not active in this channel.")

# Event listener for conversation mode translations
@bot.event
async def on_message(message):
    # Skip if the message is from the bot itself
    if message.author == bot.user:
        return
    
    # Check if the channel is in conversation mode
    if message.channel.id in conversation_channels:
        source_lang, target_lang = conversation_channels[message.channel.id]

        try:
            # Detect the language of the message to determine translation direction
            detected_lang = translator.detect(message.content).lang
            
            # Determine the translation direction based on detected language
            if detected_lang == source_lang:
                translation = translator.translate(message.content, src=source_lang, dest=target_lang)
                await message.channel.send(f"**{message.author.display_name} (translated to {target_lang}):** {translation.text}")
            elif detected_lang == target_lang:
                translation = translator.translate(message.content, src=target_lang, dest=source_lang)
                await message.channel.send(f"**{message.author.display_name} (translated to {source_lang}):** {translation.text}")
            else:
                # If the language doesn't match either, assume it’s source_lang by default
                translation = translator.translate(message.content, src=source_lang, dest=target_lang)
                await message.channel.send(f"**{message.author.display_name} (translated to {target_lang}):** {translation.text}")
                
        except Exception as e:
            print(f"Error: {e}")
            await message.channel.send("Sorry, I couldn't translate that message.")

    await bot.process_commands(message)

# Command to translate a file
@bot.tree.command(name="translate_file", description="Upload a text or docx file to translate")
@app_commands.choices(target_lang=LANGUAGE_CHOICES)
async def translate_file(interaction: discord.Interaction, target_lang: app_commands.Choice[str]):
    await interaction.response.defer()  # Defer to give the bot time to process

    temp_dir = "./temp_files"
    os.makedirs(temp_dir, exist_ok=True)  # Ensure directory exists

    # Find the latest text or docx attachment
    last_message = None
    async for message in interaction.channel.history(limit=10):
        if message.attachments and message.id != interaction.id:
            last_message = message
            break

    if not last_message or not last_message.attachments:
        await interaction.followup.send("Please upload a .txt or .docx file in the chat before running the command.")
        return

    # Download the file
    file = last_message.attachments[0]
    file_path = os.path.join(temp_dir, file.filename)
    
    if file.filename.endswith('.txt') or file.filename.endswith('.docx'):
        await file.save(file_path)

        translated_content = ""

        # Process .txt files
        if file.filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            translation = translator.translate(text, dest=target_lang.value)
            translated_content = translation.text

        # Process .docx files
        elif file.filename.endswith('.docx'):
            doc = Document(file_path)
            translated_doc = Document()
            for para in doc.paragraphs:
                if para.text.strip():
                    translation = translator.translate(para.text, dest=target_lang.value)
                    translated_doc.add_paragraph(translation.text)
                else:
                    translated_doc.add_paragraph("")
            
            translated_file_path = os.path.join(temp_dir, f"translated_{file.filename}")
            translated_doc.save(translated_file_path)
            await interaction.followup.send("Here's the translated file:", file=discord.File(translated_file_path))
            
            os.remove(file_path)
            os.remove(translated_file_path)
            return

        if translated_content:
            translated_file_path = os.path.join(temp_dir, f"translated_{file.filename}")
            with open(translated_file_path, 'w', encoding='utf-8') as f:
                f.write(translated_content)

            await interaction.followup.send("Here's the translated file:", file=discord.File(translated_file_path))
            os.remove(file_path)
            os.remove(translated_file_path)
    else:
        await interaction.followup.send("Only .txt and .docx files are supported for translation.")

# New Voice Translation Command
@bot.tree.command(name="translate_voice", description="Translate voice messages to another language")
@app_commands.choices(target_lang=LANGUAGE_CHOICES)
async def translate_voice(interaction: discord.Interaction, target_lang: app_commands.Choice[str]):
    await interaction.response.defer()  # Defer to give the bot time to process

    # Find the latest audio attachment
    last_message = None
    async for message in interaction.channel.history(limit=10):
        if message.attachments and message.id != interaction.id:
            last_message = message
            break

    if not last_message or not last_message.attachments:
        await interaction.followup.send("Please upload an audio file (.wav, .mp3, .ogg) before using the command.")
        return

    audio_file = last_message.attachments[0]
    if not audio_file.filename.endswith(('.wav', '.mp3', '.ogg')):
        await interaction.followup.send("Only .wav, .mp3, and .ogg audio files are supported.")
        return

    temp_dir = "./temp_files"
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, audio_file.filename)
    wav_path = os.path.join(temp_dir, "converted_audio.wav")  # Path for the converted file
    translated_audio_path = os.path.join(temp_dir, f"translated_{audio_file.filename}")

    # Download the audio file
    await audio_file.save(file_path)

    # Convert the audio to WAV format using pydub
    try:
        audio = AudioSegment.from_file(file_path)
        audio.export(wav_path, format="wav")

        # Convert audio to text using SpeechRecognition
        recognizer = sr.Recognizer()
        with sr.AudioFile(wav_path) as source:
            audio_data = recognizer.record(source)
            source_text = recognizer.recognize_google(audio_data, language="en")  # Adjust source language as needed

            # Translate text to target language
            translated_text = translator.translate(source_text, dest=target_lang.value).text

            # Convert translated text to speech (TTS)
            tts = gTTS(translated_text, lang=target_lang.value)
            tts.save(translated_audio_path)

            # Send the translated voice message
            await interaction.followup.send("Here's the translated voice message:", file=discord.File(translated_audio_path))

    except Exception as e:
        await interaction.followup.send(f"An error occurred: {e}")
    finally:
        # Clean up files
        for path in [file_path, wav_path, translated_audio_path]:
            if os.path.exists(path):
                os.remove(path)
@bot.tree.command(name="github", description="Get the link to the GitHub repository for the Sawti project")
async def github(interaction: discord.Interaction):
    repo_link = "https://github.com/lurantys/Sawti"
    await interaction.response.send_message(f"Check out the GitHub repository for the Sawti project here: {repo_link}")





bot.run(TOKEN)
